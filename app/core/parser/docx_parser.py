"""Word简历解析器（基于 python-docx）"""
import logging
from app.core.parser.base import BaseParser, ParseResult

logger = logging.getLogger(__name__)


class DocxParser(BaseParser):
    """Word (.docx) 文件解析器"""

    def supported_extensions(self) -> list:
        return [".docx", ".doc"]

    def parse(self, file_path: str) -> ParseResult:
        try:
            from docx import Document

            doc = Document(file_path)
            all_text = []

            # 提取段落文本（含样式信息）
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    # 如果是标题样式，添加标记
                    if para.style and para.style.name and 'Heading' in para.style.name:
                        all_text.append(f"\n## {text}")
                    else:
                        all_text.append(text)

            # 提取表格
            for table in doc.tables:
                table_text = self._extract_table(table)
                if table_text:
                    all_text.append(table_text)

            raw_text = "\n".join(all_text)
            raw_text = self._clean_text(raw_text)

            if not raw_text.strip():
                return ParseResult(
                    success=False,
                    error="Word文件中未提取到有效文本内容，文件可能为空",
                )

            logger.info(f"Word解析成功: {file_path}, {len(raw_text)}字符")
            return ParseResult(
                success=True,
                raw_text=raw_text,
                metadata={"file_path": file_path, "type": "docx"},
            )

        except ImportError:
            return ParseResult(success=False, error="python-docx库未安装")
        except Exception as e:
            logger.error(f"Word解析失败: {file_path}, 错误: {e}")
            return ParseResult(success=False, error=f"Word解析失败: {str(e)}")

    def _extract_table(self, table) -> str:
        """提取表格内容为文本"""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                rows.append(" | ".join(cells))
        return "\n".join(rows) if rows else ""

    def _clean_text(self, text: str) -> str:
        """清洗文本"""
        import re
        # 去除多余空行
        text = re.sub(r'\n{3,}', '\n\n', text)
        # 去除多余空格
        text = re.sub(r' {2,}', ' ', text)
        return text.strip()
